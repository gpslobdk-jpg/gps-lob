from __future__ import annotations

import argparse
import shutil
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph


CONTROLLER = "[UDFYLDES AF KOMMUNEN/SKOLEEJEREN]"
DOCUMENT_VERSION = "1.1"
PUBLICATION_DATE = "8. august 2026"


def clear_paragraph(paragraph) -> None:
    element = paragraph._element
    for child in list(element):
        if child.tag != qn("w:pPr"):
            element.remove(child)


def set_paragraph(paragraph, text: str, *, bold_prefix: str | None = None) -> None:
    clear_paragraph(paragraph)
    if bold_prefix and text.startswith(bold_prefix):
        prefix = paragraph.add_run(bold_prefix)
        prefix.bold = True
        text = text[len(bold_prefix) :]
    lines = text.split("\n")
    for index, line in enumerate(lines):
        if index:
            paragraph.add_run().add_break()
        paragraph.add_run(line)


def set_multiline_left(paragraph, text: str) -> None:
    set_paragraph(paragraph, text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def blank(paragraph) -> None:
    set_paragraph(paragraph, "")
    p_pr = paragraph._p.pPr
    if p_pr is not None:
        num_pr = p_pr.find(qn("w:numPr"))
        if num_pr is not None:
            p_pr.remove(num_pr)


def shade_paragraph(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, *, bold: bool = False, font_size: float = 8.0) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)


def mark_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    marker = tr_pr.find(qn("w:tblHeader"))
    if marker is None:
        marker = OxmlElement("w:tblHeader")
        tr_pr.append(marker)
    marker.set(qn("w:val"), "true")


def insert_paragraph_after(paragraph, text: str, style: str = "Normal") -> Paragraph:
    element = OxmlElement("w:p")
    paragraph._p.addnext(element)
    inserted = Paragraph(element, paragraph._parent)
    inserted.style = style
    set_paragraph(inserted, text)
    return inserted


def add_document_notice(doc: Document) -> None:
    anchor = doc.paragraphs[4]
    notice = anchor.insert_paragraph_before()
    notice.style = doc.styles["Normal"]
    notice.alignment = 1
    shade_paragraph(notice, "E8F5F0")
    run = notice.add_run(
        "STANDARD DATABEHANDLERAFTALE – IKKE UNDERSKREVET\n"
        f"Version {DOCUMENT_VERSION} · Udgivet og senest opdateret {PUBLICATION_DATE} · "
        "Dokumentansvarlig: Jeppe Laursen, SkoleGPS.dk\n"
        "Standardskabelonen skal udfyldes, vurderes og underskrives af den konkrete kommune/skoleejer "
        "og SkoleGPS. Den er ikke en myndighedsgodkendelse eller juridisk rådgivning."
    )
    run.bold = True
    run.font.color.rgb = RGBColor(6, 78, 59)
    run.font.size = Pt(10)


def fill_parties_and_choices(doc: Document) -> None:
    p = doc.paragraphs

    set_paragraph(p[9], CONTROLLER)
    set_paragraph(p[10], f"CVR {CONTROLLER}")
    set_paragraph(p[11], CONTROLLER)
    set_paragraph(p[12], CONTROLLER)
    set_paragraph(p[13], "Danmark")

    set_paragraph(p[19], "Jeppe Laursen, privatperson og ejer/driftsansvarlig for SkoleGPS.dk")
    set_paragraph(p[20], "CVR Ikke CVR-registreret (privatperson)")
    set_paragraph(p[21], "Sandbergvej 29")
    set_paragraph(p[22], "4760 Vordingborg")
    set_paragraph(p[23], "Danmark")

    set_paragraph(
        p[38],
        "I forbindelse med leveringen af SkoleGPS.dk – en webbaseret tjeneste til oprettelse og "
        "afvikling af undervisningsløb – behandler databehandleren personoplysninger på vegne af "
        "den dataansvarlige i overensstemmelse med disse Bestemmelser.",
    )

    blank(p[68])
    set_paragraph(
        p[97],
        "Databehandleren må således ikke gøre brug af en underdatabehandler til opfyldelse af disse "
        "Bestemmelser uden forudgående generel skriftlig godkendelse fra den dataansvarlige.",
    )
    set_paragraph(
        p[99],
        "Databehandleren har den dataansvarliges generelle godkendelse til brug af underdatabehandlere. "
        "Databehandleren skal skriftligt underrette den dataansvarlige om planlagte ændringer vedrørende "
        "tilføjelse eller udskiftning af underdatabehandlere med mindst 30 dages varsel og derved give "
        "den dataansvarlige mulighed for at gøre indsigelse inden brugen. Listen over godkendte "
        "underdatabehandlere fremgår af bilag B.",
    )
    blank(p[101])
    blank(p[109])

    set_paragraph(
        p[146],
        "den dataansvarliges forpligtelse til uden unødig forsinkelse og om muligt senest 72 timer, "
        "efter at denne er blevet bekendt med det, at anmelde brud på persondatasikkerheden til "
        "Datatilsynet, medmindre det er usandsynligt, at bruddet indebærer en risiko for fysiske "
        "personers rettigheder eller frihedsrettigheder",
    )
    set_paragraph(
        p[152],
        "den dataansvarliges forpligtelse til at høre Datatilsynet inden behandling, såfremt en "
        "konsekvensanalyse vedrørende databeskyttelse viser, at behandlingen vil føre til høj risiko i "
        "mangel af foranstaltninger truffet af den dataansvarlige for at begrænse risikoen.",
    )
    set_paragraph(
        p[159],
        "Databehandlerens underretning til den dataansvarlige skal om muligt ske senest 24 timer efter, "
        "at databehandleren er blevet bekendt med bruddet, sådan at den dataansvarlige kan overholde "
        "sin forpligtelse efter databeskyttelsesforordningens artikel 33.",
    )

    set_paragraph(
        p[172],
        "Ved ophør af tjenesterne vedrørende behandling af personoplysninger er databehandleren "
        "forpligtet til at slette alle personoplysninger, der er blevet behandlet på vegne af den "
        "dataansvarlige, og bekræfte over for den dataansvarlige, at oplysningerne er slettet, medmindre "
        "EU-retten eller medlemsstaternes nationale ret foreskriver opbevaring. Returnering kan i stedet "
        "aftales ved særskilt dokumenteret instruks, i det omfang det er teknisk muligt.",
    )
    for idx in (174, 176, 178):
        blank(p[idx])


def fill_signatures_and_contacts(doc: Document) -> None:
    p = doc.paragraphs
    set_paragraph(p[203], f"Navn: {CONTROLLER}")
    set_paragraph(p[204], f"Stilling: {CONTROLLER}")
    set_paragraph(p[205], f"Telefonnummer: {CONTROLLER}")
    set_paragraph(p[206], f"E-mail: {CONTROLLER}")

    set_paragraph(p[211], "Navn: Jeppe Laursen")
    set_paragraph(p[212], "Stilling: Ejer og driftsansvarlig for SkoleGPS.dk (privatperson)")
    set_paragraph(p[213], "Telefonnummer: +45 40 87 45 38")
    set_paragraph(p[214], "E-mail: skolegpsdk@gmail.com")

    set_paragraph(p[222], f"Navn: {CONTROLLER}")
    set_paragraph(p[223], f"Stilling: {CONTROLLER}")
    set_paragraph(p[224], f"Telefonnummer: {CONTROLLER}")
    set_paragraph(p[225], f"E-mail: {CONTROLLER}")

    set_paragraph(p[228], "Navn: Jeppe Laursen")
    set_paragraph(p[229], "Stilling: Ejer og driftsansvarlig for SkoleGPS.dk (privatperson)")
    set_paragraph(p[230], "Telefonnummer: +45 40 87 45 38")
    set_paragraph(p[231], "E-mail: skolegpsdk@gmail.com")


def fill_annex_a(doc: Document) -> None:
    p = doc.paragraphs
    blank(p[234])
    set_paragraph(
        p[238],
        "Formålet er at stille SkoleGPS.dk til rådighed for den dataansvarliges lærere og "
        "undervisningskonsulenter, så de kan oprette, dele, starte og gennemføre stedbaserede "
        "undervisningsløb, lade elever deltage uden elevkonto, vise opgaver og aktuelle poster, modtage "
        "besvarelser og eventuelle fotos samt give læreren status og resultater. Behandlingen må ikke "
        "anvendes til reklame, videresalg af elevdata, profilering eller andre uforenelige formål.",
    )
    set_paragraph(
        p[242],
        "Behandlingen omfatter indsamling, registrering, organisering, opbevaring, løbende opdatering, "
        "visning, begrænset videregivelse til godkendte underdatabehandlere, fejlsøgning, sletning og "
        "sikkerhedskopiering. Under et aktivt løb modtages elevens aktuelle GPS-position og overskrives "
        "løbende på deltagerposten. Positionen er ikke synlig efter 15 minutters inaktivitet og nulstilles "
        "fysisk ved næste femminutters oprydning, normalt senest efter cirka 20 minutter. Afslutning eller "
        "forladelse nulstiller straks, hvor det er muligt. Der opbygges ikke en særskilt positionshistorik for det normale "
        "løbsflow. Lærere kan dele et løb til afvikling; modtageren får en selvstændig kopi uden tidligere "
        "sessioner, deltagere, elevdata, resultater, PIN-koder eller live-status.",
    )
    set_multiline_left(
        p[246],
        "Almindelige personoplysninger:\n"
        "• Lærere/konsulenter: navn, e-mailadresse, bruger- og autentifikations-id, loginudbyder, "
        "kontostatus, oprettede undervisningsløb samt frivilligt samtykke til produktnyt.\n"
        "• Elever/deltagere: valgfrit fornavn eller holdnavn, pseudonyme sessions- og deltager-id'er, "
        "aktuel GPS-position (breddegrad, længdegrad og eventuel nøjagtighed) under aktivt løb, "
        "tidsstempler, svar, point, gennemførelsesstatus og foto ved fotoopgave.\n"
        "• Teknisk drift: IP-/netværksoplysninger hos drifts- og kortleverandører, browser, enhedstype, "
        "operativsystem, URL/sti uden hemmelige tokens, fejl- og driftsmetadata samt sikkerhedslog.\n"
        "Der skal ikke behandles CPR-numre, særlige kategorier efter artikel 9, oplysninger om strafbare "
        "forhold eller fortrolige elevsager. Genkendelige personer må ikke fotograferes.",
    )
    for idx in (248, 250, 252):
        blank(p[idx])
    set_multiline_left(
        p[256],
        "• Elever og andre deltagere i undervisningsløb, typisk børn og unge.\n"
        "• Lærere, pædagoger, undervisningskonsulenter, skoleadministratorer og andre autoriserede "
        "medarbejdere hos den dataansvarlige.\n"
        "• Kontaktpersoner hos kommunen/skoleejeren.\n"
        "Forældre/værger er kun registrerede, hvis de selv kontakter support eller indgår i en konkret, "
        "dokumenteret aktivitet.",
    )
    set_paragraph(
        p[260],
        "Behandlingen kan påbegyndes efter ikrafttræden og fortsætter, mens den dataansvarlige anvender "
        "tjenesten. De enkelte kategorier opbevares efter bilag C.4. Ved ophør slettes oplysningerne efter "
        "bestemmelse 11.1 og bilag C.4. Aftalen forudsætter gratis adgang i skoleåret 2026/27; dette "
        "ændrer ikke databeskyttelsesforpligtelserne.",
    )


def fill_annex_b(doc: Document) -> None:
    table = doc.tables[2]
    vendors = [
        (
            "Supabase, Inc.",
            "Amerikansk selskab; ikke dansk CVR",
            "Projektregion eu-west-1 (Irland) samt øvrige lokationer efter gældende DPA og underdatabehandlerliste",
            "Database, autentifikation, Storage, Realtime og serverfunktioner. Centrale konto-, løbs-, "
            "sessions-, deltager-, positions-, svar- og fotodata.",
        ),
        (
            "Vercel Inc.",
            "Delaware nr. 5857312",
            "440 N Barranca Ave #4133, Covina, CA 91723, USA; behandling efter gældende DPA og underdatabehandlerliste",
            "Webhosting, edge/serverafvikling og webanalyse. Må kun anvendes på en plan med kontraktuel "
            "DPA-dækning for kommunens personoplysninger.",
        ),
        (
            "Functional Software, Inc. (Sentry)",
            "Amerikansk selskab; ikke dansk CVR",
            "Valgt datalagringsregion Tyskland; support og øvrig behandling efter gældende DPA og underdatabehandlerliste",
            "Aktiv fejlmonitorering med applikationsredaktion. Ingen elevnavne, svar, lokationer, fotos, "
            "PIN-koder eller delingstokens må sendes.",
        ),
        (
            "OpenAI Ireland Ltd.",
            "Irsk selskab; ikke dansk CVR",
            "EØS og øvrige lokationer efter gældende DPA, overførselsgrundlag og underdatabehandlerliste",
            "Frivillige lærerrettede AI-funktioner. Elevdata, genkendelige billeder, særlige kategorier og "
            "fortroligt materiale må ikke indsendes.",
        ),
    ]

    while len(table.rows) < len(vendors) + 1:
        table.add_row()

    headers = ["NAVN", "CVR/REGISTRERING", "ADRESSE/BEHANDLINGSSTED", "BESKRIVELSE AF BEHANDLING"]
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, font_size=8.0)
    mark_table_header(table.rows[0])

    for row_index, vendor in enumerate(vendors, start=1):
        for col_index, value in enumerate(vendor):
            set_cell_text(table.rows[row_index].cells[col_index], value, font_size=7.5)

    set_paragraph(
        doc.paragraphs[273],
        "Databehandleren varsler planlagte tilføjelser eller udskiftninger mindst 30 dage før den nye "
        "underdatabehandler tages i brug. Den dataansvarlige kan inden fristens udløb gøre skriftlig og "
        "sagligt begrundet indsigelse. Bugsnag/SmartBear, Stripe og andre betingede integrationer er ikke "
        "godkendte underdatabehandlere efter denne version og må ikke aktiveres til kommunens "
        "personoplysninger uden varsling og skriftligt behandlingsgrundlag.",
    )


def fill_annex_c(doc: Document) -> None:
    p = doc.paragraphs
    set_paragraph(
        p[281],
        "Databehandleren instrueres i alene at behandle de i bilag A beskrevne oplysninger for at levere, "
        "sikre, vedligeholde og supportere SkoleGPS.dk for den dataansvarlige. Databehandleren må ikke "
        "anvende oplysningerne til egne markedsføringsformål, videresalg, reklameprofilering eller træning "
        "af offentlige AI-modeller. Den dataansvarlige instruerer sine brugere i at anvende holdnavn eller "
        "kort fornavn frem for fuldt navn, undgå særlige kategorier og undlade at fotografere genkendelige "
        "personer. AI-funktioner må ikke modtage elevdata eller andet personhenførbart materiale.",
    )
    set_paragraph(
        p[286],
        "Behandlingen vedrører navnlig børn og unge og kan under aktive løb omfatte præcis lokation, svar "
        "og fotos. Selv om særlige kategorier ikke er tilsigtet, kan et foto eller fritekstsvar utilsigtet "
        "afsløre følsomme forhold. Sikkerhedsniveauet skal derfor være højt for elevdata, og adgangen skal "
        "begrænses til den relevante lærer, den enkelte deltagers egen aktive session og strengt nødvendige "
        "driftsfunktioner. Kommunen afgør, om behandlingen kræver en konsekvensanalyse (DPIA).",
    )
    for idx in (288, 290):
        blank(p[idx])

    set_paragraph(
        p[296],
        "Elever deltager uden elevkonto og bør angive holdnavn eller kort fornavn. Interne UUID'er bruges "
        "til bruger-, sessions- og deltagerbinding. Delingslinks bruger en rå hemmelighed i URL-fragmentet, "
        "som fjernes fra browserens synlige historik efter aflæsning; alene SHA-256-digest lagres. "
        "Transport sker via HTTPS/TLS. Kryptering i hvile hos hosting-/databaseleverandører skal dokumenteres "
        "via deres aktuelle sikkerheds- og DPA-materiale.",
    )
    set_paragraph(
        p[298],
        "Adgang styres med autentifikation, row-level security, ejerskabskontrol, deltagerbinding og "
        "server-only privilegerede nøgler. Rå service-nøgler må ikke eksponeres i klienten. Medarbejder- og "
        "administratoradgang følger mindst mulige rettigheder. Der er aktuelt én driftsansvarlig; enhver "
        "senere hjælper skal have dokumenteret fortrolighed og særskilt, behovsbestemt adgang. Ændringer "
        "testes før frigivelse, og kritiske elevflows afgrænses mod utilsigtede ændringer.",
    )
    set_paragraph(
        p[300],
        "Driften baseres på hosting- og databaseleverandørernes til enhver tid gældende redundans og "
        "backupfunktioner. Der gives ikke en særskilt garanti for et bestemt RTO eller RPO i denne aftale. "
        "Databehandleren prioriterer genetablering og orienterer den dataansvarlige ved en hændelse. "
        "Leverandørernes dokumenterede backupretention gælder; Storage-objekter er ikke omfattet af "
        "Supabase-databasebackups. Kommunen kan kræve en særskilt kontinuitetsaftale før ibrugtagning.",
    )
    set_paragraph(
        p[302],
        "Der anvendes typecheck, lint, build, målrettede automatiserede tests samt fokuserede browser- og "
        "sikkerhedstests ved væsentlige ændringer. Sikkerhedsrelevante migrationer gennemgås særskilt og "
        "afprøves isoleret, før de eventuelt køres mod hosted miljø. Afhængigheder og leverandørvilkår "
        "gennemgås løbende og mindst årligt i forbindelse med DPO-/tilsynsopfølgning.",
    )
    set_paragraph(
        p[304],
        "Lærerområder kræver login. Elevadgang sker via kortvarig løbskode/QR til en aktiv session og "
        "bindes til en pseudonym deltageridentitet. Afsluttede eller ugyldige sessioner må ikke give adgang "
        "til nye data. Resultater er begrænset til løbets ejer. Delingslinks kan deaktiveres, må ikke "
        "indekseres eller caches og må ikke logges med rå token.",
    )
    set_paragraph(
        p[306],
        "Al browser-, API- og leverandørtrafik skal ske over HTTPS/TLS. Autorisationsoplysninger og "
        "service-nøgler sendes kun i relevante beskyttede kanaler. URLs, fejlrapporter og analytics skal "
        "redigeres for PIN-koder, delingstokens, navne, deltager-id'er, svar, billeder og lokation.",
    )
    set_paragraph(
        p[308],
        "Databaseadgang beskyttes med row-level security og serverkontrol. Fotoobjekter gemmes i en privat "
        "Storage-bucket. Visning kræver autentificeret lærer og kontrol af løbs-, svar- og fotoejerskab. En "
        "beskyttet, ikke-cachebar SkoleGPS-route streamer derefter billedets bytes; hverken Storage-sti eller "
        "signed Storage-URL udleveres til browseren. Uploads dekodes, rotation anvendes, og billedet genkodes "
        "som JPEG uden EXIF- eller GPS-metadata. Fotoopgaver må fortsat "
        "kun bruges til ting, steder og ikke-personhenførbare motiver. Læreren kan rydde fotos og øvrige "
        "elevdata fra resultatsiden, hvor Storage-objekt og databasepost slettes samlet.",
    )
    set_paragraph(
        p[310],
        "Primær hosting og database drives i leverandørernes datacentre med deres fysiske kontroller. "
        "Databehandlerens lokale enheder skal være adgangskodebeskyttede, opdaterede og utilgængelige for "
        "uautoriserede. Lokale downloads af elevdata må ikke foretages, medmindre den dataansvarlige har "
        "givet dokumenteret instruks og opbevaring/sletning er aftalt.",
    )
    set_paragraph(
        p[312],
        "Administrativ adgang fra hjemme-/fjernarbejdsplads må kun ske fra betroet, opdateret enhed over "
        "krypteret forbindelse og med stærk, unik adgangskode; multifaktorautentifikation skal være "
        "aktiveret, hvor leverandøren understøtter det. Skærm og session låses ved fravær, og elevdata må "
        "ikke udskrives eller deles via privat e-mail/chat.",
    )
    set_paragraph(
        p[314],
        "Der logges begrænsede drifts- og fejloplysninger. Sentry aktiveres kun ved eksplicit konfiguration, "
        "fjerner bruger-, server-, request- og enhedskontekst og redigerer navne, e-mail, IP, PIN-/løbskoder, "
        "tokens, sessions- og deltager-id'er, svar, fotostier, URL-legitimationsoplysninger og lokation. "
        "Delings- og fotoudleveringsflow er fravalgt i analytics. Vercel Analytics behandler tekniske "
        "besøgsdata. Bugsnag findes som en betinget kodeintegration, men må ikke aktiveres til kommunal "
        "behandling uden skriftlig ændringsmeddelelse, leverandøraftale og tilsvarende global redaktion. "
        "Logdata må ikke bruges til elevprofilering.",
    )

    set_paragraph(
        p[319],
        "Databehandleren bistår uden unødig forsinkelse med søgning, adgangsbegrænsning, rettelse, "
        "sletning, eksport hvor teknisk muligt, håndtering af registreredes anmodninger, vurdering af "
        "sikkerhedsbrud, dokumentation til anmeldelse, konsekvensanalyse og forudgående høring. "
        "Henvendelser fra elever/forældre henvises som udgangspunkt til den dataansvarlige.",
    )
    set_paragraph(
        p[321],
        "Læreren kan fra resultatsiden slette billeder, besvarelser, deltagere og live-sessioner for et "
        "løb. Databehandleren kan efter verificeret instruks bistå med yderligere sletning. Ved brud leveres "
        "om muligt inden for 24 timer en foreløbig beskrivelse af hændelsen, berørte kategorier, omtrentligt "
        "omfang, sandsynlige konsekvenser og trufne/foreslåede afhjælpninger; manglende oplysninger "
        "eftersendes løbende.",
    )
    set_multiline_left(
        p[324],
        "1. Fotos: private Storage-objekter og billedreferencer slettes automatisk efter 30 dage. Læreren kan "
        "slette tidligere fra resultatsiden. Hosted cron skal aktiveres og verificeres efter deployment; indtil "
        "det er dokumenteret, er lærerens manuelle oprydning den bindende driftsprocedure.\n"
        "2. Øvrige elev-/sessionsdata: svar, deltagere, navne, beskeder og afsluttede live-sessioner slettes "
        "automatisk efter 90 dage fra afslutning eller dokumenteret inaktivitet. Aktive sessioner uden "
        "retentionsanker bevares. Læreren kan slette tidligere. Hosted aktivering skal dokumenteres, før "
        "automatisk sletning oplyses som aktiv i produktion.\n"
        "3. Aktuel position: overskrives løbende og er ikke synlig efter 15 minutters inaktivitet. Den "
        "nulstilles fysisk ved næste femminutters oprydning, normalt senest efter cirka 20 minutter, og straks "
        "ved afslutning/forladelse, hvor det er muligt. Positionsfelter gemmes ikke i arkiverede svar.\n"
        "4. Lærerkonto og lærerskabte løb: opbevares, mens kontoen/aftalen er aktiv, og slettes efter "
        "dokumenteret anmodning eller ved ophør, bortset fra lovpligtig opbevaring.\n"
        "5. Tekniske oprydningslogs uden elevoplysninger slettes efter 30 dage. Øvrige driftslogs/backups "
        "følger de skriftligt godkendte leverandørvilkår; særskilte kommunale krav aftales skriftligt.\n"
        "6. Ved aftalens ophør slettes personoplysninger og eksisterende kopier i aktive systemer uden "
        "unødig forsinkelse; rester i lovlige, lukkede backups udløber efter leverandørens godkendte "
        "retention og må ikke bruges til andre formål. Databehandleren bekræfter sletningen skriftligt.",
    )
    for idx in (326, 328, 330):
        blank(p[idx])

    set_multiline_left(
        p[335],
        "• SkoleGPS' applikationsdrift: Vercel Inc.; behandling i USA og andre lokationer efter "
        "leverandørens gældende DPA, underdatabehandlerliste og overførselsgrundlag.\n"
        "• Database, autentifikation og Storage: Supabase, Inc.; den lokalt tilknyttede projektkonfiguration "
        "angiver region eu-west-1 (Irland), mens support, backups og underleverandører kan indebære andre "
        "lokationer efter leverandørens gældende DPA.\n"
        "• Fejlmonitorering: Functional Software, Inc. (Sentry); valgt datalagringsregion Tyskland, med "
        "support og underleverandører efter leverandørens gældende DPA.\n"
        "• Lærerrettet AI: OpenAI Ireland Ltd. efter gældende DPA og underdatabehandlerliste, hvis læreren "
        "frivilligt bruger funktionen; elevdata må ikke sendes.\n"
        "• Databehandlerens administrative arbejde: Sandbergvej 29, 4760 Vordingborg, Danmark.\n"
        "Andre lokaliteter eller væsentlige ændringer kræver skriftlig godkendelse efter bilag B.",
    )
    set_paragraph(
        p[338],
        "Den dataansvarlige instruerer kun til overførsel uden for EØS, hvis den konkrete leverandør, "
        "behandlingsaktivitet, lokation og overførselsmekanisme fremgår af bilag B/C og er skriftligt "
        "godkendt. Overførsler skal baseres på en gyldig afgørelse om tilstrækkeligt beskyttelsesniveau, "
        "EU-Kommissionens standardkontraktbestemmelser og nødvendige supplerende foranstaltninger eller "
        "andet gyldigt kapitel V-grundlag. Kommunen foretager/vedligeholder den nødvendige "
        "overførselsvurdering.",
    )
    set_paragraph(
        p[340],
        "For Vercel, Supabase, Sentry og OpenAI anvendes leverandørernes gældende DPA, "
        "underdatabehandlerliste og relevante overførselsgrundlag. Databehandleren må kun aktivere en "
        "leverandør til kommunens personoplysninger, når den konkrete konto og plan er omfattet af et "
        "gyldigt databehandlergrundlag. Væsentlige ændringer varsles efter bilag B.",
    )

    set_paragraph(
        p[345],
        "Den dataansvarlige kan én gang årligt og ved konkret sikkerhedshændelse eller myndighedskrav "
        "anmode om dokumentation for efterlevelse. Databehandleren leverer inden rimelig frist relevante "
        "politikker, testresultater, leverandør-DPA'er, underdatabehandlerlister og anden tilgængelig "
        "dokumentation med nødvendige fortroligheds- og sikkerhedsbegrænsninger. Hvis dokumentationen ikke "
        "er tilstrækkelig, kan den dataansvarlige gennemføre en forholdsmæssig fjernrevision eller inspektion "
        "med mindst 30 dages varsel, medmindre et brud eller myndighedskrav nødvendiggør kortere varsel. "
        "Parterne aftaler omfang og omkostninger på forhånd; databehandleren afsætter rimelig nødvendig tid.",
    )
    for idx in range(347, 371):
        if idx < len(p):
            blank(p[idx])

    set_paragraph(
        p[372],
        "Databehandleren følger mindst årligt de godkendte underdatabehandleres aktuelle sikkerheds-, "
        "revisions- og databeskyttelsesdokumentation samt væsentlige ændringer og hændelser. Tilgængelige "
        "SOC-/ISO-erklæringer, DPA'er, underdatabehandlerlister og overførselsdokumentation indhentes eller "
        "henvises til efter anmodning. Den dataansvarlige kan kræve supplerende dokumentation eller gøre "
        "indsigelse efter bilag B. Hvis en væsentlig risiko ikke kan afhjælpes, må den pågældende "
        "underdatabehandler/funktion ikke anvendes til kommunens personoplysninger.",
    )
    set_paragraph(
        p[371],
        "C.8 Procedurer for revisioner, herunder inspektioner, med behandling af personoplysninger, "
        "som er overladt til underdatabehandlere",
    )
    p[371].runs[0].bold = True
    for idx in range(374, 409):
        if idx < len(p):
            blank(p[idx])


def fill_annex_d(doc: Document) -> None:
    anchor = doc.paragraphs[411]
    sections = [
        (
            "D.1 Dokumentstatus og indgåelse",
            f"Dette dokument er Standarddatabehandleraftale – SkoleGPS, version {DOCUMENT_VERSION}, udgivet og senest "
            f"opdateret {PUBLICATION_DATE}. Den downloadede fil er en ikke underskrevet standardskabelon baseret "
            "på Datatilsynets standardbestemmelser. Kommunen/skoleejeren udfylder egne parts-, kontakt- og "
            "underskriftsfelter og foretager sin egen juridiske, sikkerhedsmæssige og eventuelle DPIA-vurdering. "
            "Dokumentet er ikke en certificering eller myndighedsgodkendelse.",
        ),
        (
            "D.2 Gratis tjeneste i skoleåret 2026/27",
            "SkoleGPS stilles gratis til rådighed for danske skoler og kommuner i skoleåret "
            "2026/27. Gratis levering ændrer ikke parternes pligter efter databeskyttelsesreglerne. "
            "Eventuelle senere betalings-, support-, oppetids-, ansvar- og opsigelsesvilkår aftales særskilt "
            "og må ikke forringe disse standardbestemmelser eller de registreredes rettigheder.",
        ),
        (
            "D.3 Obligatoriske brugsregler for den dataansvarlige",
            "Den dataansvarlige skal instruere brugerne i: (a) at anvende holdnavn eller kort fornavn frem "
            "for fuldt elevnavn, (b) ikke at indtaste CPR-numre, særlige kategorier, diagnoser, elevsager "
            "eller andre fortrolige oplysninger, (c) ikke at fotografere genkendelige personer, (d) ikke at "
            "sende elevdata eller personhenførbart materiale til AI-funktioner, (e) at rydde løbsdata efter "
            "bilag C.4 og (f) at behandle løbskoder, QR-koder og delingslinks fortroligt.",
        ),
        (
            "D.4 Foto – bindende anvendelsesbegrænsning",
            "Fotoopgaver må i kommunal brug alene omfatte ting, steder eller andre ikke-personhenførbare "
            "motiver. Genkendelige personer, elevnavne, skærmbilleder med personoplysninger og andet "
            "personhenførbart eller fortroligt indhold må ikke fotograferes eller uploades. Fotoobjekter ligger "
            "i privat Storage og streames kun til løbets ejer gennem en beskyttet, ikke-cachebar SkoleGPS-route. "
            "Storage-stier og signed Storage-URLs udleveres ikke til browseren. Ønsker kommunen senere "
            "personhenførbare fotos, kræver det en fornyet risikovurdering og skriftlig ændring af instruksen.",
        ),
        (
            "D.5 Eksterne kort-, AI- og indholdstjenester",
            "SkoleGPS' browser kan hente kortfliser/geokodning fra OpenStreetMap/Nominatim, CARTO og Esri "
            "samt enkelte offentlige medieaktiver. Lærerens frivillige AI- og indholdsværktøjer kan benytte "
            "OpenAI, Pollinations, YouTube og Apple iTunes-søgning. Disse tjenester kan modtage brugerens "
            "IP-adresse, browsermetadata og den forespurgte kortflise/søgning direkte. Der sendes ikke tilsigtet "
            "elevnavn, elevsvar, elevfoto eller elevlokation til AI- og indholdsværktøjerne. Kommunen skal "
            "instruere lærerne i kun at bruge ikke-personhenførbart materiale og kan beslutte at undlade de "
            "frivillige funktioner. Korttjenesterne modtager teknisk nødvendige IP- og forespørgselsdata.",
        ),
        (
            "D.6 Kommunens konkrete valg ved indgåelse",
            "1. Kommunens juridiske navn, CVR, adresse, kontaktperson, underskriver og behandlingsgrundlag.\n"
            "2. Kommunens vurdering af behovet for en konsekvensanalyse (DPIA) ved børn og lokation.\n"
            "3. Om fotoopgaver med ikke-personhenførbare motiver må anvendes.\n"
            "4. Om de frivillige lærerrettede AI- og indholdsfunktioner må anvendes.\n"
            "5. Kommunens ansvar for tidligere manuel sletning og kontrol af de aftalte automatiske frister.\n"
            "6. Kommunens kontaktvej ved sikkerhedsbrud og anmodninger fra registrerede.\n"
            "7. Eventuelle supplerende krav til oppetid, support, revision, ansvar, forsikring, værneting og "
            "ophør i en særskilt hovedaftale.",
        ),
    ]

    current = anchor
    for title, body in sections:
        current = insert_paragraph_after(current, title, "Normal")
        current.runs[0].bold = True
        current.runs[0].font.size = Pt(11)
        current = insert_paragraph_after(current, body)


def remove_omitted_template_scaffold(doc: Document) -> None:
    """Remove superseded examples/alternatives so empty template paragraphs do not create blank pages."""
    paragraphs = doc.paragraphs
    remove_indices = set(range(234, 236))
    remove_indices.update(range(247, 254))
    remove_indices.update(range(287, 292))
    remove_indices.update(range(325, 331))
    remove_indices.update(range(347, 371))
    remove_indices.update(range(374, 411))

    for index in sorted(remove_indices, reverse=True):
        element = paragraphs[index]._element
        element.getparent().remove(element)


def enable_field_updates(doc: Document) -> None:
    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def set_core_properties(doc: Document) -> None:
    props = doc.core_properties
    props.title = "Standarddatabehandleraftale – SkoleGPS"
    props.subject = "Datatilsynets standardkontraktbestemmelser med bilag"
    props.author = "SkoleGPS.dk"
    props.last_modified_by = "SkoleGPS.dk"
    props.comments = ""
    props.keywords = f"SkoleGPS, databehandleraftale, GDPR, standardaftale, version {DOCUMENT_VERSION}"


def patch_document_ooxml(docx_path: Path) -> None:
    """Remove template-only changelog/draft labeling and improve inherited header accessibility."""
    namespace = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    word_namespace = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    with zipfile.ZipFile(docx_path, "r") as source, tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp:
        temp_path = Path(temp.name)
        with zipfile.ZipFile(temp, "w", zipfile.ZIP_DEFLATED) as target:
            for item in source.infolist():
                payload = source.read(item.filename)
                if item.filename == "word/document.xml":
                    from lxml import etree

                    root = etree.fromstring(payload)
                    body = root.find(f"{{{word_namespace}}}body")
                    if body is not None:
                        title_block = None
                        for child in list(body):
                            if "Udkast" in "".join(child.itertext()):
                                title_block = child
                                break
                        if title_block is not None:
                            for child in list(body):
                                if child is title_block:
                                    break
                                body.remove(child)
                            body.remove(title_block)
                        annex_c_headings = root.xpath(
                            "//w:p[.//w:t[contains(., 'Bilag C')]]",
                            namespaces={"w": word_namespace},
                        )
                        for child in reversed(annex_c_headings):
                            child_text = "".join(child.itertext())
                            if "Instruks vedrørende behandling" not in child_text:
                                continue
                            previous = child.getprevious()
                            if previous is None or previous.tag != f"{{{word_namespace}}}p":
                                continue
                            page_break = any(
                                item.get(f"{{{word_namespace}}}type") == "page"
                                for item in previous.findall(f".//{{{word_namespace}}}br")
                            )
                            visible_text = "".join(previous.itertext()).strip()
                            if page_break and not visible_text:
                                body.remove(previous)
                            break
                    payload = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")
                elif item.filename.startswith("customXml/") and item.filename.endswith(".xml"):
                    payload = payload.replace(b">Udkast<", b"><")
                if item.filename.startswith("word/header") and item.filename.endswith(".xml"):
                    from lxml import etree

                    root = etree.fromstring(payload)
                    changed = False
                    for element in root.findall(f".//{{{namespace}}}docPr"):
                        if element.get("descr") or element.get("title"):
                            continue
                        name = (element.get("name") or "").lower()
                        element.set(
                            "descr",
                            "Dekorativ linje" if "connector" in name else "Datatilsynets logo",
                        )
                        changed = True
                    for row in root.findall(f".//{{{word_namespace}}}tbl/{{{word_namespace}}}tr[1]"):
                        tr_pr = row.find(f"{{{word_namespace}}}trPr")
                        if tr_pr is None:
                            tr_pr = etree.Element(f"{{{word_namespace}}}trPr")
                            row.insert(0, tr_pr)
                        marker = tr_pr.find(f"{{{word_namespace}}}tblHeader")
                        if marker is None:
                            marker = etree.SubElement(tr_pr, f"{{{word_namespace}}}tblHeader")
                        marker.set(f"{{{word_namespace}}}val", "true")
                        changed = True
                    if changed:
                        payload = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")
                target.writestr(item, payload)
    shutil.move(temp_path, docx_path)


def update_existing_document(doc: Document) -> None:
    replacements = {
        "Behandlingen omfatter indsamling,": (
            "Behandlingen omfatter indsamling, registrering, organisering, opbevaring, løbende opdatering, "
            "visning, begrænset videregivelse til godkendte underdatabehandlere, fejlsøgning, sletning og "
            "sikkerhedskopiering. Under et aktivt løb modtages elevens aktuelle GPS-position og overskrives "
            "løbende på deltagerposten. Positionen er ikke synlig efter 15 minutters inaktivitet og nulstilles "
            "fysisk ved næste femminutters oprydning, normalt senest efter cirka 20 minutter. Afslutning eller "
            "forladelse nulstiller straks, hvor det er muligt. Der opbygges ikke en særskilt positionshistorik "
            "for det normale løbsflow. Lærere kan "
            "dele et løb til afvikling; modtageren får en selvstændig kopi uden tidligere sessioner, deltagere, "
            "elevdata, resultater, PIN-koder eller live-status."
        ),
        "Databaseadgang beskyttes med row-level security": (
            "Databaseadgang beskyttes med row-level security og serverkontrol. Fotoobjekter gemmes i en privat "
            "Storage-bucket. Visning kræver autentificeret lærer og kontrol af løbs-, svar- og fotoejerskab. En "
            "beskyttet, ikke-cachebar SkoleGPS-route streamer derefter billedets bytes; hverken Storage-sti "
            "eller signed Storage-URL udleveres til browseren. Uploads dekodes, rotation anvendes, og billedet "
            "genkodes som JPEG uden EXIF- eller GPS-metadata. Fotoopgaver må "
            "fortsat kun bruges til ting, steder og ikke-personhenførbare motiver. Læreren kan rydde fotos og "
            "øvrige elevdata fra resultatsiden, hvor Storage-objekt og databasepost slettes samlet."
        ),
        "Der logges begrænsede drifts- og fejloplysninger.": (
            "Der logges begrænsede drifts- og fejloplysninger. Sentry aktiveres kun ved eksplicit konfiguration, "
            "fjerner bruger-, server-, request- og enhedskontekst og redigerer navne, e-mail, IP, PIN-/løbskoder, "
            "tokens, sessions- og deltager-id'er, svar, fotostier, URL-legitimationsoplysninger og lokation. "
            "Delings- og fotoudleveringsflow er fravalgt i analytics. Vercel Analytics behandler "
            "tekniske besøgsdata. Bugsnag findes som en betinget kodeintegration, men må ikke aktiveres til "
            "kommunal behandling uden skriftlig ændringsmeddelelse, leverandøraftale og tilsvarende global "
            "redaktion. Logdata må ikke bruges til elevprofilering."
        ),
        "1. Fotos:": (
            "1. Fotos: private Storage-objekter og billedreferencer slettes automatisk efter 30 dage. Læreren kan "
            "slette tidligere fra resultatsiden. Hosted cron skal aktiveres og verificeres efter deployment; indtil "
            "det er dokumenteret, er lærerens manuelle oprydning den bindende driftsprocedure.\n"
            "2. Øvrige elev-/sessionsdata: svar, deltagere, navne, beskeder og afsluttede live-sessioner slettes "
            "automatisk efter 90 dage fra afslutning eller dokumenteret inaktivitet. Aktive sessioner uden "
            "retentionsanker bevares. Læreren kan slette tidligere. Hosted aktivering skal dokumenteres, før "
            "automatisk sletning oplyses som aktiv i produktion.\n"
            "3. Aktuel position: overskrives løbende og er ikke synlig efter 15 minutters inaktivitet. Den "
            "nulstilles fysisk ved næste femminutters oprydning, normalt senest efter cirka 20 minutter, og "
            "straks ved afslutning/forladelse, hvor det er muligt. Positionsfelter gemmes ikke i arkiverede svar.\n"
            "4. Lærerkonto og lærerskabte løb: opbevares, mens kontoen/aftalen er aktiv, og slettes efter "
            "dokumenteret anmodning eller ved ophør, bortset fra lovpligtig opbevaring.\n"
            "5. Tekniske oprydningslogs uden elevoplysninger slettes efter 30 dage. Øvrige driftslogs/backups "
            "følger de skriftligt godkendte leverandørvilkår; særskilte kommunale krav aftales skriftligt.\n"
            "6. Ved aftalens ophør slettes personoplysninger og eksisterende kopier i aktive systemer uden "
            "unødig forsinkelse; rester i lovlige, lukkede backups udløber efter leverandørens godkendte "
            "retention og må ikke bruges til andre formål. Databehandleren bekræfter sletningen skriftligt."
        ),
        "Dette dokument er Standarddatabehandleraftale – SkoleGPS": (
            f"Dette dokument er Standarddatabehandleraftale – SkoleGPS, version {DOCUMENT_VERSION}, udgivet og "
            f"senest opdateret {PUBLICATION_DATE}. Den downloadede fil er en ikke underskrevet standardskabelon "
            "baseret på Datatilsynets standardbestemmelser. Kommunen/skoleejeren udfylder egne parts-, kontakt- "
            "og underskriftsfelter og foretager sin egen juridiske, sikkerhedsmæssige og eventuelle DPIA-vurdering. "
            "Dokumentet er ikke en certificering eller myndighedsgodkendelse."
        ),
        "Fotoopgaver må i kommunal brug": (
            "Fotoopgaver må i kommunal brug alene omfatte ting, steder eller andre ikke-personhenførbare motiver. "
            "Genkendelige personer, elevnavne, skærmbilleder med personoplysninger og andet personhenførbart eller "
            "fortroligt indhold må ikke fotograferes eller uploades. Fotoobjekter ligger i privat Storage og "
            "streames kun til løbets ejer gennem en beskyttet, ikke-cachebar SkoleGPS-route. Storage-stier og "
            "signed Storage-URLs udleveres ikke til browseren. Ønsker kommunen senere personhenførbare fotos, "
            "kræver det en fornyet risikovurdering og skriftlig ændring af instruksen."
        ),
    }

    for paragraph in doc.paragraphs:
        text = paragraph.text
        if text.startswith("STANDARD DATABEHANDLERAFTALE – IKKE UNDERSKREVET"):
            set_paragraph(
                paragraph,
                text.replace("Version 1.0", f"Version {DOCUMENT_VERSION}").replace(
                    "7. august 2026", PUBLICATION_DATE
                ),
            )
            continue

        for prefix, replacement in replacements.items():
            if text.startswith(prefix):
                set_paragraph(paragraph, replacement)
                break

        if paragraph.text.startswith("1. Kommunens juridiske navn"):
            set_paragraph(
                paragraph,
                paragraph.text.replace(
                    "5. Kommunens interne frist og ansvar for lærerens manuelle oprydning af elev-/sessionsdata.",
                    "5. Kommunens ansvar for tidligere manuel sletning og kontrol af de aftalte automatiske frister.",
                ),
            )

    set_core_properties(doc)
    enable_field_updates(doc)


def main() -> None:
    parser = argparse.ArgumentParser()
    source_group = parser.add_mutually_exclusive_group(required=True)
    source_group.add_argument("--template", type=Path)
    source_group.add_argument("--existing", type=Path)
    parser.add_argument("--output", required=True, type=Path)
    args = parser.parse_args()

    args.output.parent.mkdir(parents=True, exist_ok=True)
    source_path = args.template or args.existing
    if source_path is None:
        raise ValueError("A template or existing document is required")
    shutil.copyfile(source_path, args.output)
    doc = Document(args.output)

    if args.existing is not None:
        update_existing_document(doc)
        doc.save(args.output)
        patch_document_ooxml(args.output)
        print(args.output)
        return

    for table in doc.tables:
        if table.rows:
            mark_table_header(table.rows[0])

    # All index-based edits are completed before inserting the document notice or Annex D.
    fill_parties_and_choices(doc)
    fill_signatures_and_contacts(doc)
    fill_annex_a(doc)
    fill_annex_b(doc)
    fill_annex_c(doc)
    fill_annex_d(doc)
    remove_omitted_template_scaffold(doc)
    add_document_notice(doc)
    enable_field_updates(doc)
    set_core_properties(doc)

    doc.save(args.output)
    patch_document_ooxml(args.output)
    print(args.output)


if __name__ == "__main__":
    main()
