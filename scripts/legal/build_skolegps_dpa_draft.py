from __future__ import annotations

import argparse
import shutil
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph


CONTROLLER = "[UDFYLDES AF KOMMUNEN/SKOLEEJEREN]"
VERIFY = "[SKAL VERIFICERES FØR UNDERSKRIFT]"


def clear_paragraph(paragraph) -> None:
    element = paragraph._element
    for child in list(element):
        if child.tag != qn("w:pPr"):
            element.remove(child)


def set_paragraph(paragraph, text: str, *, bold_prefix: str | None = None) -> None:
    clear_paragraph(paragraph)
    if bold_prefix and text.startswith(bold_prefix):
        prefix = paragraph.add_run(bold_prefix)
        prefix.bold = True
        text = text[len(bold_prefix) :]
    lines = text.split("\n")
    for index, line in enumerate(lines):
        if index:
            paragraph.add_run().add_break()
        paragraph.add_run(line)


def set_multiline_left(paragraph, text: str) -> None:
    set_paragraph(paragraph, text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def blank(paragraph) -> None:
    set_paragraph(paragraph, "")
    p_pr = paragraph._p.pPr
    if p_pr is not None:
        num_pr = p_pr.find(qn("w:numPr"))
        if num_pr is not None:
            p_pr.remove(num_pr)


def shade_paragraph(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, *, bold: bool = False, font_size: float = 8.0) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)


def mark_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    marker = tr_pr.find(qn("w:tblHeader"))
    if marker is None:
        marker = OxmlElement("w:tblHeader")
        tr_pr.append(marker)
    marker.set(qn("w:val"), "true")


def insert_paragraph_after(paragraph, text: str, style: str = "Normal") -> Paragraph:
    element = OxmlElement("w:p")
    paragraph._p.addnext(element)
    inserted = Paragraph(element, paragraph._parent)
    inserted.style = style
    set_paragraph(inserted, text)
    return inserted


def add_draft_notice(doc: Document) -> None:
    anchor = doc.paragraphs[4]
    notice = anchor.insert_paragraph_before()
    notice.style = doc.styles["Normal"]
    notice.alignment = 1
    shade_paragraph(notice, "FFF3CD")
    run = notice.add_run(
        "FØRSTEUDKAST – IKKE UNDERSKREVET\n"
        "Skal gennemgås og godkendes af kommunen/skoleejeren, kommunens DPO og begge parter. "
        "Dokumentet er ikke en erklæring om, at SkoleGPS er juridisk godkendt."
    )
    run.bold = True
    run.font.color.rgb = RGBColor(114, 75, 0)
    run.font.size = Pt(10)


def fill_parties_and_choices(doc: Document) -> None:
    p = doc.paragraphs

    set_paragraph(p[9], CONTROLLER)
    set_paragraph(p[10], f"CVR {CONTROLLER}")
    set_paragraph(p[11], CONTROLLER)
    set_paragraph(p[12], CONTROLLER)
    set_paragraph(p[13], "Danmark")

    set_paragraph(p[19], "Jeppe Laursen, privatperson og ejer/driftsansvarlig for SkoleGPS.dk")
    set_paragraph(p[20], "CVR Ikke CVR-registreret (privatperson)")
    set_paragraph(p[21], "Sandbergvej 29")
    set_paragraph(p[22], "4760 Vordingborg")
    set_paragraph(p[23], "Danmark")

    set_paragraph(
        p[38],
        "I forbindelse med leveringen af SkoleGPS.dk – en webbaseret tjeneste til oprettelse og "
        "afvikling af undervisningsløb – behandler databehandleren personoplysninger på vegne af "
        "den dataansvarlige i overensstemmelse med disse Bestemmelser.",
    )

    blank(p[68])
    set_paragraph(
        p[97],
        "Databehandleren må således ikke gøre brug af en underdatabehandler til opfyldelse af disse "
        "Bestemmelser uden forudgående generel skriftlig godkendelse fra den dataansvarlige.",
    )
    set_paragraph(
        p[99],
        "Databehandleren har den dataansvarliges generelle godkendelse til brug af underdatabehandlere. "
        "Databehandleren skal skriftligt underrette den dataansvarlige om planlagte ændringer vedrørende "
        "tilføjelse eller udskiftning af underdatabehandlere med mindst 30 dages varsel og derved give "
        "den dataansvarlige mulighed for at gøre indsigelse inden brugen. Listen over godkendte "
        "underdatabehandlere fremgår af bilag B.",
    )
    blank(p[101])
    blank(p[109])

    set_paragraph(
        p[146],
        "den dataansvarliges forpligtelse til uden unødig forsinkelse og om muligt senest 72 timer, "
        "efter at denne er blevet bekendt med det, at anmelde brud på persondatasikkerheden til "
        "Datatilsynet, medmindre det er usandsynligt, at bruddet indebærer en risiko for fysiske "
        "personers rettigheder eller frihedsrettigheder",
    )
    set_paragraph(
        p[152],
        "den dataansvarliges forpligtelse til at høre Datatilsynet inden behandling, såfremt en "
        "konsekvensanalyse vedrørende databeskyttelse viser, at behandlingen vil føre til høj risiko i "
        "mangel af foranstaltninger truffet af den dataansvarlige for at begrænse risikoen.",
    )
    set_paragraph(
        p[159],
        "Databehandlerens underretning til den dataansvarlige skal om muligt ske senest 24 timer efter, "
        "at databehandleren er blevet bekendt med bruddet, sådan at den dataansvarlige kan overholde "
        "sin forpligtelse efter databeskyttelsesforordningens artikel 33.",
    )

    set_paragraph(
        p[172],
        "Ved ophør af tjenesterne vedrørende behandling af personoplysninger er databehandleren "
        "forpligtet til at slette alle personoplysninger, der er blevet behandlet på vegne af den "
        "dataansvarlige, og bekræfte over for den dataansvarlige, at oplysningerne er slettet, medmindre "
        "EU-retten eller medlemsstaternes nationale ret foreskriver opbevaring. Returnering kan i stedet "
        "aftales ved særskilt dokumenteret instruks, i det omfang det er teknisk muligt.",
    )
    for idx in (174, 176, 178):
        blank(p[idx])


def fill_signatures_and_contacts(doc: Document) -> None:
    p = doc.paragraphs
    set_paragraph(p[203], f"Navn: {CONTROLLER}")
    set_paragraph(p[204], f"Stilling: {CONTROLLER}")
    set_paragraph(p[205], f"Telefonnummer: {CONTROLLER}")
    set_paragraph(p[206], f"E-mail: {CONTROLLER}")

    set_paragraph(p[211], "Navn: Jeppe Laursen")
    set_paragraph(p[212], "Stilling: Ejer og driftsansvarlig for SkoleGPS.dk (privatperson)")
    set_paragraph(p[213], "Telefonnummer: +45 40 87 45 38")
    set_paragraph(p[214], "E-mail: skolegpsdk@gmail.com")

    set_paragraph(p[222], f"Navn: {CONTROLLER}")
    set_paragraph(p[223], f"Stilling: {CONTROLLER}")
    set_paragraph(p[224], f"Telefonnummer: {CONTROLLER}")
    set_paragraph(p[225], f"E-mail: {CONTROLLER}")

    set_paragraph(p[228], "Navn: Jeppe Laursen")
    set_paragraph(p[229], "Stilling: Ejer og driftsansvarlig for SkoleGPS.dk (privatperson)")
    set_paragraph(p[230], "Telefonnummer: +45 40 87 45 38")
    set_paragraph(p[231], "E-mail: skolegpsdk@gmail.com")


def fill_annex_a(doc: Document) -> None:
    p = doc.paragraphs
    blank(p[234])
    set_paragraph(
        p[238],
        "Formålet er at stille SkoleGPS.dk til rådighed for den dataansvarliges lærere og "
        "undervisningskonsulenter, så de kan oprette, dele, starte og gennemføre stedbaserede "
        "undervisningsløb, lade elever deltage uden elevkonto, vise opgaver og aktuelle poster, modtage "
        "besvarelser og eventuelle fotos samt give læreren status og resultater. Behandlingen må ikke "
        "anvendes til reklame, videresalg af elevdata, profilering eller andre uforenelige formål.",
    )
    set_paragraph(
        p[242],
        "Behandlingen omfatter indsamling, registrering, organisering, opbevaring, løbende opdatering, "
        "visning, begrænset videregivelse til godkendte underdatabehandlere, fejlsøgning, sletning og "
        "sikkerhedskopiering. Under et aktivt løb modtages elevens aktuelle GPS-position og overskrives "
        "løbende på deltagerposten. Der er ikke konstateret en særskilt positionshistorik for det normale "
        "løbsflow. Lærere kan dele et løb til afvikling; modtageren får en selvstændig kopi uden tidligere "
        "sessioner, deltagere, elevdata, resultater, PIN-koder eller live-status.",
    )
    set_multiline_left(
        p[246],
        "Almindelige personoplysninger:\n"
        "• Lærere/konsulenter: navn, e-mailadresse, bruger- og autentifikations-id, loginudbyder, "
        "kontostatus, oprettede undervisningsløb samt frivilligt samtykke til produktnyt.\n"
        "• Elever/deltagere: valgfrit fornavn eller holdnavn, pseudonyme sessions- og deltager-id'er, "
        "aktuel GPS-position (breddegrad, længdegrad og eventuel nøjagtighed) under aktivt løb, "
        "tidsstempler, svar, point, gennemførelsesstatus og foto ved fotoopgave.\n"
        "• Teknisk drift: IP-/netværksoplysninger hos drifts- og kortleverandører, browser, enhedstype, "
        "operativsystem, URL/sti uden hemmelige tokens, fejl- og driftsmetadata samt sikkerhedslog.\n"
        "Der skal ikke behandles CPR-numre, særlige kategorier efter artikel 9, oplysninger om strafbare "
        "forhold eller fortrolige elevsager. Genkendelige personer må ikke fotograferes.",
    )
    for idx in (248, 250, 252):
        blank(p[idx])
    set_multiline_left(
        p[256],
        "• Elever og andre deltagere i undervisningsløb, typisk børn og unge.\n"
        "• Lærere, pædagoger, undervisningskonsulenter, skoleadministratorer og andre autoriserede "
        "medarbejdere hos den dataansvarlige.\n"
        "• Kontaktpersoner hos kommunen/skoleejeren.\n"
        "Forældre/værger er kun registrerede, hvis de selv kontakter support eller indgår i en konkret, "
        "dokumenteret aktivitet.",
    )
    set_paragraph(
        p[260],
        "Behandlingen kan påbegyndes efter ikrafttræden og fortsætter, mens den dataansvarlige anvender "
        "tjenesten. De enkelte kategorier opbevares efter bilag C.4. Ved ophør slettes oplysningerne efter "
        "bestemmelse 11.1 og bilag C.4. Udkastet forudsætter gratis adgang i skoleåret 2026/27; dette "
        "ændrer ikke databeskyttelsesforpligtelserne.",
    )


def fill_annex_b(doc: Document) -> None:
    table = doc.tables[2]
    vendors = [
        (
            "Supabase, Inc.",
            "Ikke dansk CVR; kontraktenhed verificeres",
            "Valgt projektregion og øvrige lokationer efter leverandørens DPA/SCC – verificeres",
            "Database, autentifikation, Storage, Realtime og serverfunktioner. Centrale konto-, løbs-, "
            "sessions-, deltager-, positions-, svar- og fotodata.",
        ),
        (
            "Vercel Inc.",
            "Delaware nr. 5857312",
            "440 N Barranca Ave #4133, Covina, CA 91723, USA; global behandling efter DPA/SCC",
            "Webhosting, edge/serverafvikling og webanalyse. DPA-dækning for den konkrete plan skal "
            "verificeres; Vercels offentliggjorte DPA omtaler Pro/Enterprise.",
        ),
        (
            "Functional Software, Inc. (Sentry)",
            "Ikke dansk CVR",
            "Tyskland og USA samt godkendte underleverandører efter Sentrys DPA/SCC",
            "Fejlmonitorering, hvis aktiveret. Kun redigerede tekniske data; ingen elevnavne, svar, "
            "lokationer, fotos, PIN-koder eller delingstokens må sendes.",
        ),
        (
            "SmartBear Software, Inc. (Bugsnag)",
            "Ikke dansk CVR",
            "USA og godkendte underleverandører efter SmartBears DPA/SCC – konkret konto verificeres",
            "Fejlmonitorering, hvis aktiveret. Kun redigerede tekniske data; ingen særlige kategorier eller "
            "elevindhold må sendes.",
        ),
        (
            "OpenAI Ireland Ltd. / gældende OpenAI-kontraktenhed",
            "Ikke dansk CVR; kontraktenhed verificeres",
            "EØS samt eventuelle overførsler efter OpenAI DPA/SCC og aktuel underdatabehandlerliste",
            "Valgfrie lærerrettede AI-funktioner. Elevdata, genkendelige billeder, særlige kategorier og "
            "fortroligt materiale må ikke indsendes. Funktionen kan deaktiveres/udelades af instruksen.",
        ),
    ]

    while len(table.rows) < len(vendors) + 1:
        table.add_row()

    headers = ["NAVN", "CVR/REGISTRERING", "ADRESSE/BEHANDLINGSSTED", "BESKRIVELSE AF BEHANDLING"]
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, font_size=8.0)
    mark_table_header(table.rows[0])

    for row_index, vendor in enumerate(vendors, start=1):
        for col_index, value in enumerate(vendor):
            set_cell_text(table.rows[row_index].cells[col_index], value, font_size=7.5)

    set_paragraph(
        doc.paragraphs[273],
        "Databehandleren varsler planlagte tilføjelser eller udskiftninger mindst 30 dage før den nye "
        "underdatabehandler tages i brug. Den dataansvarlige kan inden fristens udløb gøre skriftlig og "
        "sagligt begrundet indsigelse. De i tabellen markerede verifikationspunkter skal være afklaret og "
        "accepteret skriftligt før underskrift eller før den pågældende funktion aktiveres.",
    )


def fill_annex_c(doc: Document) -> None:
    p = doc.paragraphs
    set_paragraph(
        p[281],
        "Databehandleren instrueres i alene at behandle de i bilag A beskrevne oplysninger for at levere, "
        "sikre, vedligeholde og supportere SkoleGPS.dk for den dataansvarlige. Databehandleren må ikke "
        "anvende oplysningerne til egne markedsføringsformål, videresalg, reklameprofilering eller træning "
        "af offentlige AI-modeller. Den dataansvarlige instruerer sine brugere i at anvende holdnavn eller "
        "kort fornavn frem for fuldt navn, undgå særlige kategorier og undlade at fotografere genkendelige "
        "personer. AI-funktioner må ikke modtage elevdata eller andet personhenførbart materiale.",
    )
    set_paragraph(
        p[286],
        "Behandlingen vedrører navnlig børn og unge og kan under aktive løb omfatte præcis lokation, svar "
        "og fotos. Selv om særlige kategorier ikke er tilsigtet, kan et foto eller fritekstsvar utilsigtet "
        "afsløre følsomme forhold. Sikkerhedsniveauet skal derfor være højt for elevdata, og adgangen skal "
        "begrænses til den relevante lærer, den enkelte deltagers egen aktive session og strengt nødvendige "
        "driftsfunktioner. Kommunen afgør, om behandlingen kræver en konsekvensanalyse (DPIA).",
    )
    for idx in (288, 290):
        blank(p[idx])

    set_paragraph(
        p[296],
        "Elever deltager uden elevkonto og bør angive holdnavn eller kort fornavn. Interne UUID'er bruges "
        "til bruger-, sessions- og deltagerbinding. Delingslinks bruger en rå hemmelighed i URL-fragmentet, "
        "som fjernes fra browserens synlige historik efter aflæsning; alene SHA-256-digest lagres. "
        "Transport sker via HTTPS/TLS. Kryptering i hvile hos hosting-/databaseleverandører skal dokumenteres "
        "via deres aktuelle sikkerheds- og DPA-materiale.",
    )
    set_paragraph(
        p[298],
        "Adgang styres med autentifikation, row-level security, ejerskabskontrol, deltagerbinding og "
        "server-only privilegerede nøgler. Rå service-nøgler må ikke eksponeres i klienten. Medarbejder- og "
        "administratoradgang følger mindst mulige rettigheder. Der er aktuelt én driftsansvarlig; enhver "
        "senere hjælper skal have dokumenteret fortrolighed og særskilt, behovsbestemt adgang. Ændringer "
        "testes før frigivelse, og kritiske elevflows afgrænses mod utilsigtede ændringer.",
    )
    set_paragraph(
        p[300],
        "Driften baseres på de godkendte hosting- og databaseleverandørers redundans og backupmuligheder. "
        "Databehandleren skal ved hændelser prioritere genetablering af tjenesten og orientere den "
        "dataansvarlige. Konkrete RTO/RPO, backupregioner og backupretention er ikke verificeret i dette "
        f"udkast og skal aftales før underskrift: {VERIFY}.",
    )
    set_paragraph(
        p[302],
        "Der anvendes typecheck, lint, build, målrettede automatiserede tests samt fokuserede browser- og "
        "sikkerhedstests ved væsentlige ændringer. Sikkerhedsrelevante migrationer gennemgås særskilt og "
        "afprøves isoleret, før de eventuelt køres mod hosted miljø. Afhængigheder og leverandørvilkår "
        "gennemgås løbende og mindst årligt i forbindelse med DPO-/tilsynsopfølgning.",
    )
    set_paragraph(
        p[304],
        "Lærerområder kræver login. Elevadgang sker via kortvarig løbskode/QR til en aktiv session og "
        "bindes til en pseudonym deltageridentitet. Afsluttede eller ugyldige sessioner må ikke give adgang "
        "til nye data. Resultater er begrænset til løbets ejer. Delingslinks kan deaktiveres, må ikke "
        "indekseres eller caches og må ikke logges med rå token.",
    )
    set_paragraph(
        p[306],
        "Al browser-, API- og leverandørtrafik skal ske over HTTPS/TLS. Autorisationsoplysninger og "
        "service-nøgler sendes kun i relevante beskyttede kanaler. URLs, fejlrapporter og analytics skal "
        "redigeres for PIN-koder, delingstokens, navne, deltager-id'er, svar, billeder og lokation.",
    )
    set_paragraph(
        p[308],
        "Databaseadgang beskyttes med row-level security og serverkontrol. Fotoobjekter slettes automatisk "
        "efter 30 dage og kan slettes tidligere manuelt. Den aktuelle Storage-bucket participant-uploads er dog "
        "konfigureret som offentlig i migrationskoden. Det er en væsentlig afvigelse: inden kommunen "
        "anvender fotoopgaver med personoplysninger, skal bucket ændres til privat adgang med kortlivede "
        "signerede links eller kommunen skal udtrykkeligt acceptere en anden dokumenteret model. Indtil da "
        "må genkendelige personer ikke fotograferes.",
    )
    set_paragraph(
        p[310],
        "Primær hosting og database drives i leverandørernes datacentre med deres fysiske kontroller. "
        "Databehandlerens lokale enheder skal være adgangskodebeskyttede, opdaterede og utilgængelige for "
        "uautoriserede. Lokale downloads af elevdata må ikke foretages, medmindre den dataansvarlige har "
        "givet dokumenteret instruks og opbevaring/sletning er aftalt.",
    )
    set_paragraph(
        p[312],
        "Administrativ adgang fra hjemme-/fjernarbejdsplads må kun ske fra betroet, opdateret enhed over "
        "krypteret forbindelse og med stærk, unik adgangskode; multifaktorautentifikation skal være "
        "aktiveret, hvor leverandøren understøtter det. Skærm og session låses ved fravær, og elevdata må "
        "ikke udskrives eller deles via privat e-mail/chat.",
    )
    set_paragraph(
        p[314],
        "Der logges begrænsede drifts- og fejloplysninger. Applikationen har redaktion af navne, e-mail, "
        "PIN-/løbskoder, tokens, sessions- og deltager-id'er, svar, fotos og lokation før observability. "
        "Delingssiden er særskilt fravalgt i analytics/fejlmonitorering. Aktuel retention og aktivering hos "
        f"Vercel Analytics, Sentry, Bugsnag og database-logning skal verificeres: {VERIFY}. Logdata må ikke "
        "bruges til elevprofilering.",
    )

    set_paragraph(
        p[319],
        "Databehandleren bistår uden unødig forsinkelse med søgning, adgangsbegrænsning, rettelse, "
        "sletning, eksport hvor teknisk muligt, håndtering af registreredes anmodninger, vurdering af "
        "sikkerhedsbrud, dokumentation til anmeldelse, konsekvensanalyse og forudgående høring. "
        "Henvendelser fra elever/forældre henvises som udgangspunkt til den dataansvarlige.",
    )
    set_paragraph(
        p[321],
        "Læreren kan fra resultatsiden slette billeder, besvarelser, deltagere og live-sessioner for et "
        "løb. Databehandleren kan efter verificeret instruks bistå med yderligere sletning. Ved brud leveres "
        "om muligt inden for 24 timer en foreløbig beskrivelse af hændelsen, berørte kategorier, omtrentligt "
        "omfang, sandsynlige konsekvenser og trufne/foreslåede afhjælpninger; manglende oplysninger "
        "eftersendes løbende.",
    )
    set_multiline_left(
        p[324],
        "1. Fotos: automatisk sletning af Storage-objekt og nulstilling af billedlink efter 30 dage; tidligere "
        "manuel sletning er mulig. Det hostede oprydningsjobs faktiske aktivering skal verificeres.\n"
        "2. Øvrige elev-/sessionsdata: den dataansvarlige instruerer lærerne i at rydde data senest 30 dage "
        "efter aktiviteten, medmindre et kortere dokumenteret formål gælder. Koden dokumenterer aktuelt "
        "manuel, ikke universel automatisk tidsbaseret, sletning.\n"
        "3. Aktuel position: overskrives løbende og slettes sammen med deltager-/sessionsdata.\n"
        "4. Lærerkonto og lærerskabte løb: opbevares, mens kontoen/aftalen er aktiv, og slettes efter "
        "dokumenteret anmodning eller ved ophør, bortset fra lovpligtig opbevaring.\n"
        "5. Driftslogs/backups: leverandørspecifik retention skal oplyses og godkendes før underskrift.\n"
        "6. Ved aftalens ophør slettes personoplysninger og eksisterende kopier i aktive systemer uden "
        "unødig forsinkelse; rester i lovlige, lukkede backups udløber efter leverandørens godkendte "
        "retention og må ikke bruges til andre formål. Databehandleren bekræfter sletningen skriftligt.",
    )
    for idx in (326, 328, 330):
        blank(p[idx])

    set_multiline_left(
        p[335],
        "• SkoleGPS' applikationsdrift: Vercel Inc.; primær behandling i USA og globalt efter "
        "leverandørens DPA/SCC.\n"
        "• Database, autentifikation og Storage: Supabase, Inc.; den valgte projektregion og supplerende "
        f"support-/backup-lokationer skal dokumenteres: {VERIFY}.\n"
        "• Fejlmonitorering: Sentry (Tyskland/USA) og Bugsnag/SmartBear (USA og underleverandører), kun "
        "hvis aktiveret.\n"
        "• Lærerrettet AI: OpenAI efter gældende DPA/SCC, hvis aktiveret; ingen elevdata må sendes.\n"
        "• Databehandlerens administrative arbejde: Sandbergvej 29, 4760 Vordingborg, Danmark.\n"
        "Andre lokaliteter eller væsentlige ændringer kræver skriftlig godkendelse efter bilag B.",
    )
    set_paragraph(
        p[338],
        "Den dataansvarlige instruerer kun til overførsel uden for EØS, hvis den konkrete leverandør, "
        "behandlingsaktivitet, lokation og overførselsmekanisme fremgår af bilag B/C og er skriftligt "
        "godkendt. Overførsler skal baseres på en gyldig afgørelse om tilstrækkeligt beskyttelsesniveau, "
        "EU-Kommissionens standardkontraktbestemmelser og nødvendige supplerende foranstaltninger eller "
        "andet gyldigt kapitel V-grundlag. Kommunen foretager/vedligeholder den nødvendige "
        "overførselsvurdering.",
    )
    set_paragraph(
        p[340],
        "For Vercel, Supabase, Sentry, SmartBear/Bugsnag og OpenAI skal kontraktenhed, DPA, "
        "underdatabehandlerliste, lokationer, eventuel EU-U.S. Data Privacy Framework-certificering og/eller "
        f"relevante SCC-moduler dokumenteres før underskrift. Status: {VERIFY}.",
    )

    set_paragraph(
        p[345],
        "Den dataansvarlige kan én gang årligt og ved konkret sikkerhedshændelse eller myndighedskrav "
        "anmode om dokumentation for efterlevelse. Databehandleren leverer inden rimelig frist relevante "
        "politikker, testresultater, leverandør-DPA'er, underdatabehandlerlister og anden tilgængelig "
        "dokumentation med nødvendige fortroligheds- og sikkerhedsbegrænsninger. Hvis dokumentationen ikke "
        "er tilstrækkelig, kan den dataansvarlige gennemføre en forholdsmæssig fjernrevision eller inspektion "
        "med mindst 30 dages varsel, medmindre et brud eller myndighedskrav nødvendiggør kortere varsel. "
        "Parterne aftaler omfang og omkostninger på forhånd; databehandleren afsætter rimelig nødvendig tid.",
    )
    for idx in range(347, 371):
        if idx < len(p):
            blank(p[idx])

    set_paragraph(
        p[372],
        "Databehandleren følger mindst årligt de godkendte underdatabehandleres aktuelle sikkerheds-, "
        "revisions- og databeskyttelsesdokumentation samt væsentlige ændringer og hændelser. Tilgængelige "
        "SOC-/ISO-erklæringer, DPA'er, underdatabehandlerlister og overførselsdokumentation indhentes eller "
        "henvises til efter anmodning. Den dataansvarlige kan kræve supplerende dokumentation eller gøre "
        "indsigelse efter bilag B. Hvis en væsentlig risiko ikke kan afhjælpes, må den pågældende "
        "underdatabehandler/funktion ikke anvendes til kommunens personoplysninger.",
    )
    set_paragraph(
        p[371],
        "C.8 Procedurer for revisioner, herunder inspektioner, med behandling af personoplysninger, "
        "som er overladt til underdatabehandlere",
    )
    p[371].runs[0].bold = True
    for idx in range(374, 409):
        if idx < len(p):
            blank(p[idx])


def fill_annex_d(doc: Document) -> None:
    anchor = doc.paragraphs[411]
    sections = [
        (
            "D.1 Udkastets status og forudsætninger",
            "Dette dokument er et grundigt førsteudkast baseret på Datatilsynets standardbestemmelser og "
            "en lokal teknisk kodegennemgang pr. 7. august 2026. Det er ikke underskrevet, er ikke juridisk "
            "rådgivning og udgør ikke en certificering eller myndighedsgodkendelse. Kommunen/skoleejeren "
            "skal lade egen DPO og relevante IT-/sikkerhedsfunktioner gennemgå aftalen og de markerede "
            "verifikationspunkter før brug eller underskrift.",
        ),
        (
            "D.2 Gratis tjeneste i skoleåret 2026/27",
            "SkoleGPS forventes stillet gratis til rådighed for danske skoler og kommuner i skoleåret "
            "2026/27. Gratis levering ændrer ikke parternes pligter efter databeskyttelsesreglerne. "
            "Eventuelle senere betalings-, support-, oppetids-, ansvar- og opsigelsesvilkår aftales særskilt "
            "og må ikke forringe disse standardbestemmelser eller de registreredes rettigheder.",
        ),
        (
            "D.3 Obligatoriske brugsregler for den dataansvarlige",
            "Den dataansvarlige skal instruere brugerne i: (a) at anvende holdnavn eller kort fornavn frem "
            "for fuldt elevnavn, (b) ikke at indtaste CPR-numre, særlige kategorier, diagnoser, elevsager "
            "eller andre fortrolige oplysninger, (c) ikke at fotografere genkendelige personer, (d) ikke at "
            "sende elevdata eller personhenførbart materiale til AI-funktioner, (e) at rydde løbsdata efter "
            "bilag C.4 og (f) at behandle løbskoder, QR-koder og delingslinks fortroligt.",
        ),
        (
            "D.4 Foto – midlertidig sikkerhedsgate",
            "Fotoopgaver må i kommunal brug alene omfatte ting, steder eller andre ikke-personhenførbare "
            "motiver. Hvis kommunen ønsker fotos med personer eller andre personoplysninger, er funktionen "
            "ikke godkendt efter dette udkast, før Storage er ændret til privat adgang med passende "
            "kortlivede links, ændringen er testet og dokumenteret, og kommunen har foretaget en fornyet "
            "risikovurdering/DPO-godkendelse.",
        ),
        (
            "D.5 Eksterne kort- og indholdstjenester",
            "SkoleGPS' browser kan hente kortfliser/geokodning fra OpenStreetMap/Nominatim, CARTO og Esri "
            "samt enkelte offentlige medieaktiver. Disse tjenester kan modtage brugerens IP-adresse, "
            "browsermetadata og den forespurgte kortflise/søgning direkte. Der sendes ikke tilsigtet "
            "elevnavn, svar eller foto. Tjenesternes rolle, vilkår, behandlingssteder og kapitel V-grundlag "
            "er ikke endeligt klassificeret i dette udkast. Kommunen skal godkende dem, eller de skal "
            "erstattes/deaktiveres for kommunens brug.",
        ),
        (
            "D.6 Punkter der skal lukkes før underskrift",
            "1. Kommunens juridiske navn, CVR, adresse, kontaktperson, underskriver og behandlingsgrundlag.\n"
            "2. Kommunens vurdering af behovet for DPIA ved behandling af børn og lokation.\n"
            "3. Privat adgangsmodel for foto-Storage eller dokumenteret fravalg af personhenførbare fotos.\n"
            "4. Supabase-projektregion, backup-/logretention, DPA/SCC og aktuel hosted cleanup-job.\n"
            "5. Vercel-planens kontraktuelle DPA-dækning og overførselsvurdering.\n"
            "6. Om Sentry, Bugsnag og Vercel Analytics er aktive, samt deres retention og region.\n"
            "7. Kortleverandørernes rolle og godkendelse.\n"
            "8. AI-funktionernes afgrænsning, DPA/SCC eller deaktivering.\n"
            "9. Om kommunen kræver automatisk sletning af alle øvrige elevdata frem for dokumenteret "
            "læreroprensning.\n"
            "10. Revisionsniveau, ansvar, forsikring, værneting, support og ophør i en særskilt hovedaftale.",
        ),
    ]

    current = anchor
    for title, body in sections:
        current = insert_paragraph_after(current, title, "Normal")
        current.runs[0].bold = True
        current.runs[0].font.size = Pt(11)
        current = insert_paragraph_after(current, body)


def remove_omitted_template_scaffold(doc: Document) -> None:
    """Remove superseded examples/alternatives so empty template paragraphs do not create blank pages."""
    paragraphs = doc.paragraphs
    remove_indices = set(range(234, 236))
    remove_indices.update(range(247, 254))
    remove_indices.update(range(287, 292))
    remove_indices.update(range(325, 331))
    remove_indices.update(range(347, 371))
    remove_indices.update(range(374, 411))

    for index in sorted(remove_indices, reverse=True):
        element = paragraphs[index]._element
        element.getparent().remove(element)


def enable_field_updates(doc: Document) -> None:
    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def set_core_properties(doc: Document) -> None:
    props = doc.core_properties
    props.title = "SkoleGPS.dk – databehandleraftale (førsteudkast)"
    props.subject = "Datatilsynets standardkontraktbestemmelser med bilag"
    props.author = "SkoleGPS.dk"
    props.last_modified_by = "SkoleGPS.dk"
    props.comments = "Førsteudkast. Skal gennemgås af kommunen/skoleejerens DPO før underskrift."
    props.keywords = "SkoleGPS, databehandleraftale, GDPR, førsteudkast"


def patch_header_accessibility(docx_path: Path) -> None:
    """Add useful alt text to the two header shapes inherited from the official template."""
    namespace = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    word_namespace = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    with zipfile.ZipFile(docx_path, "r") as source, tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp:
        temp_path = Path(temp.name)
        with zipfile.ZipFile(temp, "w", zipfile.ZIP_DEFLATED) as target:
            for item in source.infolist():
                payload = source.read(item.filename)
                if item.filename.startswith("word/header") and item.filename.endswith(".xml"):
                    from lxml import etree

                    root = etree.fromstring(payload)
                    changed = False
                    for element in root.findall(f".//{{{namespace}}}docPr"):
                        if element.get("descr") or element.get("title"):
                            continue
                        name = (element.get("name") or "").lower()
                        element.set(
                            "descr",
                            "Dekorativ linje" if "connector" in name else "Datatilsynets logo",
                        )
                        changed = True
                    for row in root.findall(f".//{{{word_namespace}}}tbl/{{{word_namespace}}}tr[1]"):
                        tr_pr = row.find(f"{{{word_namespace}}}trPr")
                        if tr_pr is None:
                            tr_pr = etree.Element(f"{{{word_namespace}}}trPr")
                            row.insert(0, tr_pr)
                        marker = tr_pr.find(f"{{{word_namespace}}}tblHeader")
                        if marker is None:
                            marker = etree.SubElement(tr_pr, f"{{{word_namespace}}}tblHeader")
                        marker.set(f"{{{word_namespace}}}val", "true")
                        changed = True
                    if changed:
                        payload = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")
                target.writestr(item, payload)
    shutil.move(temp_path, docx_path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--template", required=True, type=Path)
    parser.add_argument("--output", required=True, type=Path)
    args = parser.parse_args()

    args.output.parent.mkdir(parents=True, exist_ok=True)
    shutil.copyfile(args.template, args.output)
    doc = Document(args.output)

    for table in doc.tables:
        if table.rows:
            mark_table_header(table.rows[0])

    # All index-based edits are completed before inserting the draft notice or Annex D.
    fill_parties_and_choices(doc)
    fill_signatures_and_contacts(doc)
    fill_annex_a(doc)
    fill_annex_b(doc)
    fill_annex_c(doc)
    fill_annex_d(doc)
    remove_omitted_template_scaffold(doc)
    add_draft_notice(doc)
    enable_field_updates(doc)
    set_core_properties(doc)

    doc.save(args.output)
    patch_header_accessibility(args.output)
    print(args.output)


if __name__ == "__main__":
    main()
